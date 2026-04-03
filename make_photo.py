# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def sf(run, name, size, bold=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def h(doc, text, font='黑体', size=16):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, font, size, True)

def para(doc, text, font='宋体', size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p2 = doc.add_paragraph(); p2.alignment = align
    p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
    r = p2.add_run(text); sf(r, font, size)

def bc(cell, color='CCCCCC'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top','left','bottom','right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); tcBorders.append(e)
    tcPr.append(tcBorders)

def sc(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''; p2 = cell.paragraphs[0]; p2.alignment = align
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text); sf(r, '宋体', size, bold)

# 封面
para(doc, '', '宋体', 8)
for _ in range(3): para(doc, '', '宋体', 6)
h(doc, '《XX氏宗谱》', '标宋', 26)
para(doc, '图文混排版', '宋体', 12)
for _ in range(4): para(doc, '', '宋体', 6)
h(doc, '卷　首', '黑体', 14)
para(doc, '图文并茂　每人一页　可插入照片', '宋体', 11)
for _ in range(4): para(doc, '', '宋体', 6)
h(doc, '公元二零二六年三月　续修', '宋体', 12)
doc.add_page_break()

h(doc, '一、图文混排版说明', '黑体', 14)
para(doc, '图文混排版是现代家谱的重要创新，在传统文字记载的基础上，为每位族人预留照片位置，使族谱更加生动真实。每一页分为上下两部分：上部为照片区，下部为文字记载区。', '宋体', 11)

h(doc, '二、版面结构', '黑体', 14)
para(doc, '每人一页，或每代一页（视照片数量而定）。上半部分为照片区（约占页面40%），下半部分为族人基本信息及生平文字记载。', '宋体', 11)

h(doc, '三、人物卡片示例（一人一页）', '黑体', 14)

# 照片区
photo_para = doc.add_paragraph()
photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
photo_para.paragraph_format.space_before = Pt(6)
photo_para.paragraph_format.space_after = Pt(6)
r = photo_para.add_run('[ 照 片 位 置 ]\n请在此处插入族人照片\n建议尺寸：宽8cm × 高10cm（3:4比例）\n分辨率：300万像素以上\n格式：jpg 或 png')
sf(r, '宋体', 12)

# 分隔线
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(6)
sep.paragraph_format.space_after = Pt(6)
sep.add_run('─' * 40)

# 文字信息
h(doc, '基本信息', '黑体', 12)
info_table = doc.add_table(rows=3, cols=4)
info_table.alignment = 1
info_table.style = 'Table Grid'
info_items = [
    ['世次：第一世', '姓名：XXX', '字：XX', '号：XX'],
    ['生年：1900年', '卒年：1980年', '寿：80岁', '学历：XX大学'],
    ['配偶：配XXX', '子：2人', '女：1人', '职业：XX'],
]
widths_info = [4.5, 4.5, 4.5, 4.5]
for ri, row_items in enumerate(info_items):
    for ci, text in enumerate(row_items):
        cell = info_table.rows[ri].cells[ci]
        sc(cell, text, 10, bold=(ci == 0))
        bc(cell)
        cell.width = Cm(widths_info[ci])

para(doc, '', '宋体', 4)
h(doc, '生平事迹', '黑体', 12)
bio = doc.add_paragraph()
bio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
bio.paragraph_format.space_before = Pt(4)
bio.paragraph_format.space_after = Pt(4)
r = bio.add_run('简历：XXX于清光绪年间出生于XX省XX县，自幼聪颖，毕业于XX大学，曾任XX职务。公为人谦和，乐善好施，热心族务，为乡里所敬重。育有二子一女，子孙繁衍，家道兴盛。卒后葬于XX山，子孙后代至今祭祀不绝。')
sf(r, '宋体', 11)

para(doc, '', '宋体', 8)
h(doc, '四、使用说明', '黑体', 14)
for num, s in [
    ('一', '将族人照片扫描或拍照，保存为jpg格式，建议分辨率不低于300万像素。'),
    ('二', '将照片插入上方的照片区，调整至合适大小。'),
    ('三', '在下方的文字区填写族人基本信息及生平事迹。'),
    ('四', '每页可放置一人，也可一代一页，视资料多寡灵活安排。'),
]:
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(4)
    r1 = p3.add_run(f'（{num}）{s}'); sf(r1, '宋体', 11)

doc.save('C:/Users/Administrator/Downloads/jiapu_photo.docx')
print("OK")
