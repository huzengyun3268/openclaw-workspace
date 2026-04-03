# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(3)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def set_font(run, name='宋体', size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, font='宋体', size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, font, size, True)
    return p

def add_p(doc, text, font='宋体', size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, font, size)
    return p

def add_line(doc, text):
    """竖排单行"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, '宋体', 11)
    return p

# 封面
for _ in range(4): add_p(doc, '')
add_h(doc, '《×氏族谱》', '标宋', 26)
add_p(doc, '传统竖排版', '宋体', 12)
for _ in range(5): add_p(doc, '')
add_h(doc, '卷　首', '黑体', 14)
add_p(doc, '竖排右翻　繁体从右至左', '宋体', 11)
for _ in range(5): add_p(doc, '')
add_h(doc, '公元二零二六年三月　续修', '宋体', 12)

doc.add_page_break()

add_h(doc, '一、传统族谱说明', '黑体', 14)
add_p(doc, '本谱采用传统竖排版式，自右向左翻阅，繁体字印刷，以存古制。竖排族谱历史悠久，自宋代欧、苏二式定型以来，为各姓氏所通用。', '宋体', 11)
add_p(doc, '本谱分为"世系图"与"行传"两大部分，世系图标明血脉传承，行传详述个人生平，二者互为经纬，缺一不可。', '宋体', 11)

add_h(doc, '二、竖排版特点', '黑体', 14)
rules = [
    ('一', '本谱自右向左翻页，右端为卷首，左端为卷尾。'),
    ('二', '文字竖排，每行从上至下，字序从右至左。'),
    ('三', '世系图以竖线相连，代际分明，长幼有序。'),
    ('四', '行传按代数分列，每人独立成段。'),
    ('五', '采用繁体字印刷，以存古籍形制。'),
    ('六', '生卒纪年以农历为主，括注公历。'),
]
for num, r in rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'（{num}）{r}')
    set_font(r1, '宋体', 11)

add_h(doc, '三、称谓规范', '黑体', 14)
terms = [
    '讳——男子名讳，书于卷首。',
    '字——男子字号，书于讳后。',
    '号——别号，有则书之。',
    '配——正室妻子，书"配×氏"。',
    '娶——续娶，书"娶×氏"。',
    '适——女子出嫁，书"适×家"。',
    '赘——男子入赘，书"赘于×家"。',
    '卒——去世，书"卒于×年×月×日"。',
    '葬——安葬，书"葬于×山向"。',
]
for t in terms:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(t)
    set_font(r1, '宋体', 11)

add_h(doc, '四、行传格式示例', '黑体', 14)
sample_lines = [
    '第一世　父：×××　讳：×××　字：××',
    '生卒：公元一九零零年三月十五日（庚子年）生',
    '　　　公元一九八零年十月初五（庚申年）卒',
    '配　配×氏，×××之女，生卒年月日',
    '　　　生子女：×子×女',
    '简历：毕业于××大学，曾任××，热心公益',
]
for line in sample_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(line)
    set_font(r1, '宋体', 11)

add_h(doc, '五、迁徙录', '黑体', 14)
migration_lines = [
    '始迁祖×××，于清光绪年间（约公元一八八零年），',
    '自原籍××省××县××村，迁居现址。',
    '迁徙原因：战乱╱灾荒╱谋生╱官职（据实填写）。',
    '迄今已传×代，计×口丁。',
]
for line in migration_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(line)
    set_font(r1, '宋体', 11)

add_h(doc, '六、家训八则', '黑体', 14)
家训 = [
    '一、孝亲敬长，敦邻睦族',
    '二、耕读传家，勤劳致富',
    '三、诚实守信，戒赌戒毒',
    '四、遵纪守法，爱国爱乡',
    '五、男女平等，婚丧从简',
    '六、讲究卫生，健身强体',
    '七、赈灾救荒，扶困济贫',
    '八、修桥筑路，热心公益',
]
for j in 家训:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(j)
    set_font(r1, '宋体', 11)

doc.save('C:/Users/Administrator/Downloads/家谱模板_传统竖排版.docx')
print("OK - 传统竖排版")
