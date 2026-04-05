from docx import Document
import os, sys, re

sys.stdout.reconfigure(encoding='utf-8')

files = [
    (r'E:\胡氏宗谱改正版\南岙胡宗谱序-已清理.txt', '南岙胡宗谱序'),
    (r'E:\胡氏宗谱改正版\南岙胡胡氏宗谱谱丁简介-已清理.txt', '南岙胡胡氏宗谱谱丁简介'),
    (r'E:\胡氏宗谱改正版\总谱世系-已清理.txt', '总谱世系'),
]

for txt_path, name in files:
    print(f'Converting: {name}')
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Skip header lines
        lines = content.split('\n')
        body_start = 0
        for i, line in enumerate(lines):
            if i < 2:  # skip "=== title ===" and "共 X chars"
                continue
            if '===' in line:
                body_start = i + 1
                break
        
        body = '\n'.join(lines[body_start:])
        # Remove duplicate large blocks
        body_lines = [l.strip() for l in body.split('\n') if l.strip() and len(l.strip()) > 4]
        
        doc = Document()
        doc.core_properties.title = name
        title = doc.add_heading(name, 0)
        title.alignment = 1
        
        # Split into paragraphs by double newlines
        paragraphs = re.split(r'\n{2,}', body)
        for para in paragraphs:
            para = para.strip()
            if para:
                try:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.space_after = Pt(6)
                except:
                    pass
        
        out_path = rf'E:\胡氏宗谱改正版\{name}-Word.docx'
        doc.save(out_path)
        print(f'  -> {out_path}')
        print(f'  Preview: {body[:100].strip()}...')
        
    except Exception as e:
        print(f'Error: {e}')

print('\nDone!')
