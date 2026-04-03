# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 页面设置：A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)

# 辅助函数：设置中文字体
def set_run_font(run, font_name='宋体', font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置东亚字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_para(doc, text, font_name='宋体', font_size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name, font_size, bold)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 22, 2: 16, 3: 14}
    fonts = {1: '黑体', 2: '黑体', 3: '黑体'}
    add_para(doc, text, fonts[level], sizes[level], True, WD_ALIGN_PARAGRAPH.CENTER, 12, 6)

def set_cell_text(cell, text, font_name='宋体', font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    # 设置单元格边框
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ========== 封面 ==========
add_para(doc, '', font_size=8, space_before=60)
add_heading(doc, '《 ×× 氏 族 谱》', 1)
add_para(doc, '', font_size=6, space_before=6)
add_para(doc, '（公元二零二六年续修）', '宋体', 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 40)

add_para(doc, '', font_size=8, space_before=80)
add_para(doc, '编纂：×××', '宋体', 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 8)
add_para(doc, '校对：×××', '宋体', 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 8)
add_para(doc, '校对：×××', '宋体', 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 20)
add_para(doc, '公元二零二六年三月', '宋体', 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0)

doc.add_page_break()

# ========== 谱序 ==========
add_heading(doc, '一、谱 序', 2)
add_para(doc, '    ×××氏族谱续修，历时×月而成。吾族自始祖×××迁居此地，迄今已传×世，历时约×百年。先祖耕读传家，积德行善，子孙繁衍，人丁日盛。今续修族谱，以正世系、明昭穆、传家风、启后人。', '宋体', 11, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 6, 12)

# ========== 凡例 ==========
add_heading(doc, '二、凡 例', 2)
rules = [
    ('一', '本谱以横排欧式为主，世代横行，父子竖线相连，同世并列。'),
    ('二', '本谱纪年采用公历，括注农历干支，如：2026年3月27日（丙午年二月廿八）。'),
    ('三', '本谱所载人名，男子书讳，女子书氏，出嫁书"适"，招赘书"赘"。'),
    ('四', '生卒年月日时皆备，不详者书"不详"，未娶者书"未娶"，未嫁者书"待字"。'),
    ('五', '本谱所列配偶，已卒者书"卒"，在世者不注。'),
    ('六', '本谱女性与男性同等入谱，子女均注其名。'),
    ('七', '本谱所载人名，有字、号者皆附注，无者从略。'),
    ('八', '本谱资料来源于各户填报，并经逐户核实，如有错漏，望后人更正。'),
]
for num, rule in rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f'（{num}）{rule}')
    set_run_font(r1, '宋体', 11, False)

doc.add_page_break()

# ========== 世系图 ==========
add_heading(doc, '三、世 系 图', 2)
add_para(doc, '（五世一表，世代横行，父子竖线相连）', '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 10)

# 创建世系表（五服示例）
def add_lineage_table(doc, generations=6):
    """生成简化世系表"""
    for gen in range(1, generations + 1):
        add_para(doc, f'第{gen}世', '黑体', 11, True, WD_ALIGN_PARAGRAPH.LEFT, 10, 4)
        
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # 设置列宽
        for cell in table.columns[0].cells:
            cell.width = Cm(3.2)
        for cell in table.columns[1].cells:
            cell.width = Cm(3.2)
        for cell in table.columns[2].cells:
            cell.width = Cm(3.2)
        for cell in table.columns[3].cells:
            cell.width = Cm(3.2)
        for cell in table.columns[4].cells:
            cell.width = Cm(3.2)
        
        row = table.rows[0]
        names = ['一世  二子', '一世  长子', '一世  次子', '一世  三子', '一世  四子']
        for i, cell in enumerate(row.cells):
            set_cell_text(cell, names[i], '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER)
        
        doc.add_paragraph()

add_lineage_table(doc, 5)

doc.add_page_break()

# ========== 行传格式示例 ==========
add_heading(doc, '四、行 传（齿 录）', 2)
add_para(doc, '（每人一行，详细记录生平事迹）', '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER, 0, 10)

# 行传表头
table2 = doc.add_table(rows=2, cols=7)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'

# 列宽
col_widths = [1.5, 2.2, 2.0, 4.5, 3.0, 2.5, 2.8]
headers = ['世次', '姓名', '字号', '生卒年月', '配偶', '子女', '简历/学历/职业']

# 表头行
for i, (h, w) in enumerate(zip(headers, col_widths)):
    cell = table2.rows[0].cells[i]
    cell.width = Cm(w)
    set_cell_text(cell, h, '黑体', 9, True, WD_ALIGN_PARAGRAPH.CENTER)

# 示例数据行
sample = ['一世', 'XX', 'XX', '1900年3月1日（庚子年）-1980年10月15日（庚申年）', '配XX氏', '子：XX XX 女：XX', '毕业于XX大学，曾任XX']
for i, (text, w) in enumerate(zip(sample, col_widths)):
    cell = table2.rows[1].cells[i]
    cell.width = Cm(w)
    set_cell_text(cell, text, '宋体', 9, False, WD_ALIGN_PARAGRAPH.CENTER)

# 再加几行空白的
for _ in range(5):
    row = table2.add_row()
    for i, w in enumerate(col_widths):
        cell = row.cells[i]
        cell.width = Cm(w)
        set_cell_text(cell, '', '宋体', 9, False, WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ========== 迁徙录 ==========
add_heading(doc, '五、迁 徙 录', 2)
add_para(doc, '    本族始祖×××，于清光绪年间（约1880年）自原籍××省××县××村迁居现址，迄今已历×代，×口丁。', '宋体', 11, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 6, 12)
add_para(doc, '    迁徙原因：因战乱/灾荒/谋生/官职/婚姻等（详述）', '宋体', 11, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0, 12)

doc.add_paragraph()
add_para(doc, '    世系图：', '黑体', 11, True, WD_ALIGN_PARAGRAPH.LEFT, 6, 4)
migration_table = doc.add_table(rows=3, cols=4)
migration_table.alignment = WD_TABLE_ALIGNMENT.CENTER
migration_table.style = 'Table Grid'
mh_headers = ['世代', '时间', '迁徙地', '原因']
mh_widths = [2.0, 3.0, 4.0, 5.5]
for i, (h, w) in enumerate(zip(mh_headers, mh_widths)):
    cell = migration_table.rows[0].cells[i]
    cell.width = Cm(w)
    set_cell_text(cell, h, '黑体', 10, True, WD_ALIGN_PARAGRAPH.CENTER)
mh_data = [
    ['一世', '清光绪年间', '××省××县', '灾荒'],
    ['×世', '××年', '××省××县', '谋生'],
]
for ri, row_data in enumerate(mh_data):
    for ci, text in enumerate(row_data):
        cell = migration_table.rows[ri+1].cells[ci]
        cell.width = Cm(mh_widths[ci])
        set_cell_text(cell, text, '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ========== 家训族规 ==========
add_heading(doc, '六、家 训 族 规', 2)
rules2 = [
    '一、孝亲敬长，敦邻睦族',
    '二、耕读传家，勤劳致富',
    '三、诚实守信，戒赌戒毒',
    '四、遵纪守法，爱国爱乡',
    '五、男女平等，婚丧从简',
    '六、讲究卫生，健身强体',
    '七、赈灾救荒，扶困济贫',
    '八、修桥筑路，热心公益',
]
for r in rules2:
    add_para(doc, r, '宋体', 11, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 3, 3)

doc.add_paragraph()

# ========== 大事记 ==========
add_heading(doc, '七、大 事 记', 2)
events = [
    ('1900年', '始祖×××迁居本地'),
    ('1920年', '建祠堂一座'),
    ('1950年', '完成第一次续谱'),
    ('1980年', '重修祠堂'),
    ('2000年', '完成第二次续谱'),
    ('2026年', '完成第三次续谱，本谱电子化'),
]
events_table = doc.add_table(rows=len(events)+1, cols=2)
events_table.alignment = WD_TABLE_ALIGNMENT.CENTER
events_table.style = 'Table Grid'
set_cell_text(events_table.rows[0].cells[0], '年份', '黑体', 10, True, WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(events_table.rows[0].cells[1], '事件', '黑体', 10, True, WD_ALIGN_PARAGRAPH.CENTER)
events_table.rows[0].cells[0].width = Cm(3)
events_table.rows[0].cells[1].width = Cm(12)
for i, (year, event) in enumerate(events):
    set_cell_text(events_table.rows[i+1].cells[0], year, '宋体', 10, False, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(events_table.rows[i+1].cells[1], event, '宋体', 10, False, WD_ALIGN_PARAGRAPH.LEFT)
    events_table.rows[i+1].cells[0].width = Cm(3)
    events_table.rows[i+1].cells[1].width = Cm(12)

doc.add_paragraph()

# ========== 后记 ==========
add_heading(doc, '八、后 记', 2)
add_para(doc, '    族谱之修，上慰祖宗在天之灵，下启子孙继往之志。本次续修，得族人踊跃出资、鼎力相助，编委会成员不辞辛劳、反复核实，历时×月乃成。望后人珍惜此谱，每逢续修之期，踊跃参与，使我族历史文化薪火相传、永续不绝焉。', '宋体', 11, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 6, 20)
add_para(doc, '    编委会成员：×××（主编）、×××（副主编）、×××、×××、×××', '宋体', 11, False, WD_ALIGN_PARAGRAPH.LEFT, 0, 8)
add_para(doc, '    公元二零二六年三月', '宋体', 11, False, WD_ALIGN_PARAGRAPH.RIGHT, 0, 0)

# 保存
output_path = 'C:/Users/Administrator/Downloads/家谱模板_现代通用版.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
