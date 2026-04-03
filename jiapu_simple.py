# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)  # 横排A4
section.page_height = Cm(21)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def sf(run, name='宋体', size=9, bold=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def h(doc, text, font='黑体', size=14):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, font, size, True); return p

def p(doc, text, font='宋体', size=9):
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text); sf(r, font, size); return p2

def bc(cell, color='BBBBBB'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top','left','bottom','right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); tcBorders.append(e)
    tcPr.append(tcBorders)

def sc(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''; p2 = cell.paragraphs[0]; p2.alignment = align
    p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
    r = p2.add_run(text); sf(r, '宋体', size, bold)

# 封面
p(doc, '', 8)
for _ in range(3): p(doc, '', 4)
h(doc, '《×氏宗谱》', '标宋', 24)
p(doc, '现代简洁版 · 表格体', '宋体', 11)
for _ in range(4): p(doc, '', 6)
h(doc, '公元二零二六年三月　续修', '宋体', 11)
doc.add_page_break()

h(doc, '一、简洁版说明', '黑体', 13)
p(doc, '现代简洁版采用紧凑的表格形式，在一页A4纸上可容纳较多信息。适合信息量较大、需要批量录入的家族使用。', '宋体', 9)

# 超大表格
h(doc, '二、世系总表', '黑体', 13)

cols = 12
table = doc.add_table(rows=21, cols=cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表头
headers = ['世', '姓名', '字', '号', '生年', '卒年', '寿', '配偶', '子', '女', '学历/职业', '备注']
widths = [0.9, 2.0, 1.5, 1.5, 1.8, 1.8, 0.8, 2.5, 1.5, 1.5, 3.0, 3.5]
for ci, (hdr, w) in enumerate(zip(headers, widths)):
    cell = table.rows[0].cells[ci]
    sc(cell, hdr, 8, True)
    cell.width = Cm(w); bc(cell)

# 数据行
sample_data = ['一', 'XX', 'XX', '', '1900', '1980', '80', '配XX氏', '2', '1', '大学', '']
for row in table.rows[1:]:
    for ci, (text, w) in enumerate(zip(sample_data, widths)):
        cell = row.cells[ci]
        sc(cell, text, 8)
        cell.width = Cm(w); bc(cell)

p(doc, '', 4)
p(doc, f'注：共{20}行空格，可自行复制行。寿=卒年-生年。', '宋体', 8)

doc.add_page_break()

h(doc, '三、简洁版使用说明', '黑体', 13)
rules = [
    ('一', '本表为通用格式，可直接打印使用。'),
    ('二', '每行填一人，复制行可增加行数。'),
    ('三', '学历/职业栏填写最高学历及主要职业。'),
    ('四', '备注栏填写其他需要说明的事项。'),
    ('五', '如需增加列，可在Word中右键表格-插入列。'),
    ('六', '此版本适合电脑录入后打印，也可手写。'),
]
for num, r in rules:
    p2 = doc.add_paragraph(); p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
    r1 = p2.add_run(f'（{num}）{r}'); sf(r1, '宋体', 9)

doc.save('C:/Users/Administrator/Downloads/家谱模板_简洁版.docx')
print("OK - 简洁版")
