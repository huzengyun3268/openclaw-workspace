import olefile, os, sys, re
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

files = [
    (r'E:\胡氏宗谱改正版\南岙胡宗谱序.p65', '南岙胡宗谱序'),
    (r'E:\胡氏宗谱改正版\南岙胡胡氏宗谱谱丁简介.p65', '南岙胡胡氏宗谱谱丁简介'),
    (r'E:\胡氏宗谱改正版\总谱世系.p65', '总谱世系')
]

def extract_clean_text(data):
    """Extract Chinese text from PageMaker stream, removing font garbage"""
    text = data.decode('gbk', errors='ignore')
    
    # Strategy: find all valid Chinese character sequences and join with newlines
    # Pattern: sequences of Chinese chars (4+ in a row)
    chinese_seq = re.findall(r'[\u4e00-\u9fff]{4,}', text)
    
    # Also find sequences of Chinese + numbers mixed
    # Pattern: Chinese chars and numbers separated by common punctuation
    lines = []
    # Split by known separators that appear between text blocks
    # Font name garbage often looks like: "方正姚体" "华文新魏" etc
    
    # Better approach: find sections that have high density of Chinese chars
    chunks = text.split('\r')
    for chunk in chunks:
        chinese_count = sum(1 for c in chunk if 0x4e00 <= ord(c) <= 0x9fff)
        total_len = len(chunk.strip())
        if total_len > 0 and chinese_count / total_len > 0.3:
            # Clean this chunk
            cleaned = re.sub(r'[\u0000-\u007f]+', ' ', chunk)  # remove ASCII
            cleaned = ''.join(c for c in cleaned if 0x4e00 <= ord(c) <= 0x9fff or c in '，。、；：！？""''（）【】《》—…· \n')
            cleaned = re.sub(r' {2,}', ' ', cleaned)
            cleaned = cleaned.strip()
            if len(cleaned) > 5:
                lines.append(cleaned)
    
    return '\n'.join(lines)

for filepath, name in files:
    print(f'Processing: {name}...')
    try:
        ole = olefile.OleFileIO(filepath)
        data = ole.openstream('PageMaker').read()
        ole.close()
        
        text = extract_clean_text(data)
        
        # Save clean txt
        txt_path = rf'E:\胡氏宗谱改正版\{name}-清晰文字版.txt'
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f'=== {name} ===\n共提取 {len(text)} 个字符\n\n')
            f.write(text)
        print(f'  TXT: {len(text)} chars -> {txt_path}')
        
        # Save as DOCX
        doc = Document()
        doc.core_properties.title = name
        title = doc.add_heading(name, 0)
        title.alignment = 1
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc_path = rf'E:\胡氏宗谱改正版\{name}-清晰文字版.docx'
        doc.save(doc_path)
        print(f'  DOCX saved')
        
    except Exception as e:
        print(f'Error: {e}')

print('\nDone!')
