# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def sf(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def h(doc, text, font='黑体', size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    sf(r, font, size, True)
    return p

def para(doc, text, font='宋体', size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p2 = doc.add_paragraph()
    p2.alignment = align
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(3)
    r = p2.add_run(text)
    sf(r, font, size)
    return p2

def line(doc, text, font='宋体', size=11):
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(2)
    p3.paragraph_format.first_line_indent = Pt(22)
    r = p3.add_run(text)
    sf(r, font, size)
    return p3

# 封面
para(doc, '', font='宋体', size=8)
for _ in range(3):
    para(doc, '', font='宋体', size=6)
h(doc, '《XX氏宗谱》', '标宋', 26)
para(doc, '牒记体版', '宋体', 12)
for _ in range(4):
    para(doc, '', font='宋体', size=6)
h(doc, '卷　首', '黑体', 14)
para(doc, '牒记体　纯文字叙述　每人一段　按代数编撰', '宋体', 11)
for _ in range(4):
    para(doc, '', font='宋体', size=6)
h(doc, '公元二零二六年三月　续修', '宋体', 12)
doc.add_page_break()

h(doc, '一、牒记体说明', '黑体', 14)
para(doc, '牒记体，又称"纪传体"，是传统族谱的重要体例之一。每人独立成段，详细记载其生平事迹、迁徙经历、职业成就、道德品行等。文字可长可短，视资料多寡而定。', '宋体', 11)
para(doc, '牒记体的优点：资料详尽，可读性强；有利于保存家族史料；便于后人了解先祖生平。缺点：录入工作量较大；需要较多的家族史料支撑。', '宋体', 11)

h(doc, '二、编写要求', '黑体', 14)
for num, r in [
    ('一', '每人一节，独立成段，首行注明世次与姓名。'),
    ('二', '生卒年月、籍贯、迁徙情况须准确记载。'),
    ('三', '配偶、子女简述，重点记述其特长、成就。'),
    ('四', '有功名、官职、学历、技艺者，从实记载。'),
    ('五', '女性与男性同等对待，记其姓名及适何家。'),
    ('六', '无法考证者，注明"不详"，不臆造。'),
]:
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(4)
    r1 = p3.add_run(f'（{num}）{r}')
    sf(r1, '宋体', 11)

h(doc, '三、牒记体行传示例', '黑体', 14)
p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(6)
p3.paragraph_format.space_after = Pt(4)
r1 = p3.add_run('【第一世】')
sf(r1, '黑体', 12)

for ex in [
    '始祖　讳：XXX　字：XX　号：XX',
    '生于清光绪XX年XX月XX日（农历），卒于XX年XX月XX日（农历），享年XX岁。',
    '原籍XX省XX县XX村，清光绪年间迁居现址。',
    '配X氏，XXX之女，生卒不详，葬于XX山向。',
    '生子女X人：子X人、女X人，俱各已婚配。',
    '公少有大志，耕读传家，为人谦和，乐善好施。',
    '卒后葬于XX，乡里至今称颂其德。',
]:
    line(doc, ex)

para(doc, '', '宋体', 6)
p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(6)
p3.paragraph_format.space_after = Pt(4)
r1 = p3.add_run('【第二世】')
sf(r1, '黑体', 12)

for ex in [
    '二世祖　讳：XXX　父：始祖XXX',
    '生于民国XX年XX月XX日，卒于公元XXXX年XX月XX日，享年XX岁。',
    '配X氏，XXX之女，生于XX年，卒于XX年。',
    '毕业于XX大学，曾任XX职务，热心族务。',
    '生子女X人：子：XXX、XXX；女：XXX、XXX。',
]:
    line(doc, ex)

doc.add_page_break()
h(doc, '四、迁徙录', '黑体', 14)
para(doc, '本族迁徙历程，据各房家传及老谱记载，整理如下：', '宋体', 11)
line(doc, '始迁祖XXX，约于清光绪十年（公元一八八四年），因战乱频仍、民生艰难，自原籍XX省XX县XX村，携家带口，历尽艰辛，迁至现址定居。初来时仅有茅屋三间，薄田数亩，艰苦创业。')
line(doc, '二世祖XXX，承父业，勤耕作，家道渐兴。至三世祖时，人口繁衍，遂成村落。')
line(doc, '此后数代，或因谋生、或因从军、或因婚姻，族人散居各地，今已遍布XX省XX县及周边县市。')

h(doc, '五、大事记', '黑体', 14)
for ev in [
    '清光绪十年（1884）——始祖迁居本地',
    '清宣统三年（1911）——辛亥革命，族人参与反正',
    '民国二十年（1931）——修谱一次',
    '公元1950年——土地改革',
    '公元1980年——改革开放，族人外出务工经商',
    '公元2000年——第二次续谱',
    '公元2026年——第三次续修，增补牒记体行传',
]:
    line(doc, ev)

doc.save('C:/Users/Administrator/Downloads/jiapu_diejil.docx')
print("OK")
