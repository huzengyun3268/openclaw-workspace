import olefile, os, sys, re
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding='utf-8')

# Known Chinese font names that appear in PageMaker files
FONT_GARBAGE = [
    '方正姚体', '方正宋体', '方正黑体', '方正楷体', '方正舒体',
    '华文新魏', '华文行楷', '华文隶书', '华文幼圆', '华文琥珀',
    '华文彩云', '华文仿宋', '华文黑体', '华文宋体', '华文细黑',
    '文鼎CS', '文鼎简', '文鼎GB',
    '汉仪报', '汉仪宋', '汉仪黑', '汉仪楷',
    '金山WPS', '金山简', '金山宋', '金山黑',
    '微软雅黑', '微软黑体', '微软宋体', '微软仿宋',
    '黑体', '宋体', '楷体', '仿宋', '隶书', '幼圆',
    '综艺体', '琥珀体', '彩云体', '细黑体',
    '新魏', '行楷', '舒体', '姚体',
    '思源黑体', '思源宋体',
    '文泉驿',
    '文档主页', '默认',
]

# Build regex patterns
font_patterns = []
for f in FONT_GARBAGE:
    if f not in font_patterns:
        font_patterns.append(re.escape(f))

# Also strip sequences of ASCII and numbers that appear between Chinese text
# Common patterns: 878, 88, 0, 4, 2 etc at the start of lines

def clean_text(text):
    """Remove font garbage and clean text"""
    # Step 1: Remove all font name occurrences
    for fp in font_patterns:
        text = re.sub(fp, '', text)
    
    # Step 2: Remove ASCII digits and some punctuation repeated
    # Remove standalone numbers like 888, 34, 0, 4, etc that appear isolated
    text = re.sub(r'^[0-9\s]{1,20}$', '', text, flags=re.MULTILINE)  # lines with only numbers
    text = re.sub(r'\b[0-9]{1,10}\b', ' ', text)  # isolated numbers
    text = re.sub(r'[0-9]+\s*[0-9]+', ' ', text)  # number groups
    
    # Step 3: Remove special chars
    text = re.sub(r'[▲▼●○◆◇□■□▪▫]', '', text)
    
    # Step 4: Remove very short lines (likely garbage)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # Count Chinese chars
        chinese = sum(1 for c in line if 0x4e00 <= ord(c) <= 0x9fff)
        total = len(line.strip())
        if total == 0:
            continue
        ratio = chinese / total if total > 0 else 0
        # Keep lines with good Chinese ratio and reasonable length
        if ratio > 0.5 and len(line.strip()) >= 4:
            cleaned_lines.append(line.strip())
    
    return '\n'.join(cleaned_lines)

def extract_text_from_stream(data):
    """Extract text from PageMaker binary stream"""
    text = data.decode('gbk', errors='ignore')
    
    # Remove null bytes
    text = text.replace('\x00', '')
    
    # Split into lines by carriage return
    lines = text.split('\r')
    
    result = []
    for line in lines:
        # Remove null chars
        line = line.strip()
        if not line:
            continue
        # Count Chinese
        chinese = sum(1 for c in line if 0x4e00 <= ord(c) <= 0x9fff)
        if chinese >= 3:  # at least 3 Chinese chars
            result.append(line)
    
    return '\n'.join(result)

files = [
    (r'E:\胡氏宗谱改正版\南岙胡宗谱序.p65', '南岙胡宗谱序'),
    (r'E:\胡氏宗谱改正版\南岙胡胡氏宗谱谱丁简介.p65', '南岙胡胡氏宗谱谱丁简介'),
    (r'E:\胡氏宗谱改正版\总谱世系.p65', '总谱世系')
]

for filepath, name in files:
    print(f'Processing: {name}...')
    try:
        ole = olefile.OleFileIO(filepath)
        data = ole.openstream('PageMaker').read()
        ole.close()
        
        # Extract and clean
        raw_text = extract_text_from_stream(data)
        cleaned = clean_text(raw_text)
        
        # Remove duplicate consecutive lines
        lines = cleaned.split('\n')
        deduped = []
        prev = ''
        for line in lines:
            line = line.strip()
            if line and line != prev and len(line) > 2:
                deduped.append(line)
                prev = line
        
        final_text = '\n'.join(deduped)
        
        print(f'  Raw text: {len(raw_text)} chars -> Clean: {len(final_text)} chars')
        
        # Save clean TXT
        txt_path = rf'E:\胡氏宗谱改正版\{name}-已清理.txt'
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f'=== {name} ===\n共 {len(final_text)} 个字符\n\n')
            f.write(final_text)
        print(f'  -> {txt_path}')
        
        # Save as DOCX
        doc = Document()
        doc.core_properties.title = name
        title = doc.add_heading(name, 0)
        title.alignment = 1
        
        for line in final_text.split('\n'):
            line = line.strip()
            if line:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(3)
        
        doc_path = rf'E:\胡氏宗谱改正版\{name}-已清理.docx'
        doc.save(doc_path)
        print(f'  -> {doc_path}')
        print(f'  Preview: {final_text[:100]}...')
        
    except Exception as e:
        print(f'Error: {e}')
        import traceback
        traceback.print_exc()

print('\nDone!')
