import olefile, os, sys, re
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding='utf-8')

files = [
    (r'E:\胡氏宗谱改正版\南岙胡宗谱序.p65', '南岙胡宗谱序'),
    (r'E:\胡氏宗谱改正版\南岙胡胡氏宗谱谱丁简介.p65', '南岙胡胡氏宗谱谱丁简介'),
    (r'E:\胡氏宗谱改正版\总谱世系.p65', '总谱世系')
]

FONT_GARBAGE = [
    '方正姚体', '方正宋体', '方正黑体', '方正楷体', '方正舒体',
    '华文新魏', '华文行楷', '华文隶书', '华文幼圆', '华文琥珀',
    '华文彩云', '华文仿宋', '华文黑体', '华文宋体', '华文细黑',
    '文鼎CS', '文鼎简', '文鼎GB',
    '汉仪报', '汉仪宋', '汉仪黑', '汉仪楷',
    '金山WPS', '金山简', '金山宋', '金山黑',
    '微软雅黑', '微软黑体', '微软宋体', '微软仿宋',
    '黑体', '宋体', '楷体', '仿宋', '隶书', '幼圆',
    '综艺体', '琥珀体', '彩云体', '细黑体',
    '新魏', '行楷', '舒体', '姚体',
    '文泉驿',
    '文档主页', '默认',
]

def clean_for_docx(text):
    """Clean text for docx compatibility"""
    # Remove control chars except newline/tab
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Remove font garbage
    for f in FONT_GARBAGE:
        text = text.replace(f, '')
    # Remove isolated numbers/garbage lines
    text = re.sub(r'^[0-9\s]{1,30}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\b[0-9]{1,15}\b', ' ', text)
    return text.strip()

def extract_text_from_stream(data):
    text = data.decode('gbk', errors='ignore')
    text = text.replace('\x00', '')
    lines = text.split('\r')
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        chinese = sum(1 for c in line if 0x4e00 <= ord(c) <= 0x9fff)
        if chinese >= 3:
            result.append(line)
    return '\n'.join(result)

for filepath, name in files:
    print(f'Processing: {name}...')
    try:
        ole = olefile.OleFileIO(filepath)
        data = ole.openstream('PageMaker').read()
        ole.close()
        
        raw_text = extract_text_from_stream(data)
        cleaned = clean_for_docx(raw_text)
        
        # Remove duplicates
        lines = cleaned.split('\n')
        deduped = []
        prev = ''
        for line in lines:
            line = line.strip()
            if line and line != prev and len(line) > 2:
                deduped.append(line)
                prev = line
        
        final = '\n'.join(deduped)
        print(f'  Text length: {len(final)} chars')
        
        # Save as DOCX
        doc = Document()
        doc.core_properties.title = name
        
        # Title
        title = doc.add_heading(name, 0)
        title.alignment = 1
        
        # Paragraphs
        for line in final.split('\n'):
            line = line.strip()
            if line:
                try:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(3)
                except:
                    pass
        
        doc_path = rf'E:\胡氏宗谱改正版\{name}.docx'
        doc.save(doc_path)
        print(f'  -> {doc_path}')
        
        # Show preview
        print(f'  Preview: {final[:80]}...')
        
    except Exception as e:
        print(f'Error: {e}')
        import traceback
        traceback.print_exc()

print('\nAll done!')
