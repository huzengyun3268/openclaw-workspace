# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def sf(run, name='宋体', size=11, bold=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def h(doc, text, font='黑体', size=16):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, font, size, True); return p

def p(doc, text, font='宋体', size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p2 = doc.add_paragraph(); p2.alignment = align
    p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
    r = p2.add_run(text); sf(r, font, size); return p2

def bc(cell, color='CCCCCC'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top','left','bottom','right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); tcBorders.append(e)
    tcPr.append(tcBorders)

def sc(cell, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''; p2 = cell.paragraphs[0]; p2.alignment = align
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text); sf(r, '宋体', size, bold)

# 封面
p(doc, '', 8)
for _ in range(3): p(doc, '', 6)
h(doc, '《×氏宗谱》', '标宋', 26)
p(doc, '图文混排版', '宋体', 12)
for _ in range(4): p(doc, '', 6)
h(doc, '卷　首', '黑体', 14)
p(doc, '图文并茂　每人一页　可插入照片', '宋体', 11)
for _ in range(4): p(doc, '', 6)
h(doc, '公元二零二六年三月　续修', '宋体', 12)
doc.add_page_break()

h(doc, '一、图文混排版说明', '黑体', 14)
p(doc, '图文混排版是现代家谱的重要创新，在传统文字记载的基础上，为每位族人预留照片位置，使族谱更加生动真实。每一页分为上下两部分：上部为照片区，下部为文字记载区。', '宋体', 11)

h(doc, '二、版面结构', '黑体', 14)
p(doc, '每人一页，或每代一页（视照片数量而定）。上半部分预留照片位置（约占页面40%），下半部分为族人基本信息及生平文字记载。', '宋体', 11)

h(doc, '三、图文人物卡（示例）', '黑体', 14)

# 人物卡片表格
card = doc.add_table(rows=2, cols=3)
card.alignment = WD_TABLE_ALIGNMENT.CENTER
card.style = 'Table Grid'

# 照片区（合并3列）
photo_cell = card.rows[0].cells[0]
# 合并单元格
photo_cell.merge(card.rows[0].cells[1])
photo_cell.merge(card.rows[0].cells[2])
photo_cell.text = ''
pp = photo_cell.paragraphs[0]
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pp.paragraph_format.space_before = Pt(40)
pp.paragraph_format.space_after = Pt(40)
r = pp.add_run('[此处插入照片]\n建议尺寸：3:4 比例')
sf(r, '宋体', 10)
for b in ['top','left','bottom','right']:
    pass  # 不加边框
# 文字区
info = [
    ['世次：第一世', '姓名：XX', '字号：XX'],
    ['生卒：1900年-1980年', '配偶：配XX氏', '子女：XX'],
    ['简历：毕业于XX大学，曾任XX职务。\n热心族务，敦亲睦族，为乡里所敬重。', '', ''],
]
info_row = card.rows[1].cells
# 第一行
for ci, text in enumerate(['世次：第一世', '姓名：XX', '字号：XX']):
    sc(info_row[ci], text, 10, False)
    bc(info_row[ci])

doc.add_paragraph()

# 再加一行信息
t2 = doc.add_table(rows=1, cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.style = 'Table Grid'
for ci, text in enumerate(['生卒：1900年-1980年', '配偶：配XX氏', '子女：XX']):
    sc(t2.rows[0].cells[ci], text, 10)
    bc(t2.rows[0].cells[ci])

t3 = doc.add_table(rows=1, cols=1)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3.style = 'Table Grid'
t3.rows[0].cells[0].text = ''
pp3 = t3.rows[0].cells[0].paragraphs[0]
pp3.paragraph_format.space_before = Pt(4); pp3.paragraph_format.space_after = Pt(4)
r3 = pp3.add_run('简历：毕业于XX大学，曾任XX职务。热心族务，为乡里所敬重。')
sf(r3, '宋体', 10)
bc(t3.rows[0].cells[0])

p(doc, '', 6)
p(doc, '注：照片建议尺寸3:4，300万像素以上，jpg格式。文字区可容纳约300字生平简介。', '宋体', 9)

h(doc, '四、使用说明', '黑体', 14)
steps = [
    ('一', '将族人照片扫描或拍照，保存为jpg格式，建议分辨率不低于300万像素。'),
    ('二', '将照片插入上方的照片区，调整至合适大小。'),
    ('三', '在下方的文字区填写族人基本信息及生平事迹。'),
    ('四', '每页可放置一人，也可一代一页，视资料多寡灵活安排。'),
]
for num, s in steps:
    p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(4)
    r1 = p3.add_run(f'（{num}）{s}'); sf(r1, '宋体', 11)

doc.save('C:/Users/Administrator/Downloads/家谱模板_图文混排版.docx')
print("OK - 图文混排版")
