# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def set_font(run, name='宋体', size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, font='黑体', size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, font, size, True)
    return p

def add_p(doc, text, font='宋体', size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, font, size)
    return p

def set_cell(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, '宋体', size, bold)

def border_cell(cell, color='999999'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        e = OxmlElement(f'w:{b}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)

# 封面
add_p(doc, '', size=8)
for _ in range(4): add_p(doc, '', size=4)
add_h(doc, '《×氏宗谱》', '标宋', 28)
add_p(doc, '苏式垂珠体版', '宋体', 12)
for _ in range(5): add_p(doc, '', size=6)
add_h(doc, '卷　首', '黑体', 14)
add_p(doc, '垂珠体世系图　竖排右翻　直系传承', '宋体', 11)
for _ in range(5): add_p(doc, '', size=8)
add_h(doc, '公元二零二六年三月　续修', '宋体', 12)

doc.add_page_break()

# 苏式说明
add_h(doc, '一、苏式垂珠体说明', '黑体', 14)
add_p(doc, '苏式世系，又称"垂珠体"，其特点是代际以竖线相连，形如垂珠。直系居中，旁支分列两侧，层次清晰。其优点在于：直系血脉一目了然；旁支关系分明；适合支系不复杂的中小家族。', '宋体', 11)
add_p(doc, '本谱采用苏式排版，自始祖起，以竖线连接各代，同辈横向排列，直系居中，旁支分列左右，以线条粗细区分主支。', '宋体', 11)

add_h(doc, '二、阅读方法', '黑体', 14)
rules = [
    ('一', '本谱自右向左翻阅，始祖居最右端。'),
    ('二', '直系传承居中，以粗线连接上下代。'),
    ('三', '旁支分列左右，以细线与直系相连。'),
    ('四', '同辈横向排列，依长幼自左向右。'),
    ('五', '每人名下注生卒、配偶、子女概要。'),
]
for num, r in rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'（{num}）{r}')
    set_font(r1, '宋体', 11)

add_h(doc, '三、垂珠体世系表示例', '黑体', 14)

# 竖排世系表
# 用表格展示垂珠效果
table = doc.add_table(rows=5, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 世代标签
gen_labels = ['第一世', '第二世', '第三世', '第四世', '第五世']
for i, label in enumerate(gen_labels):
    cell = table.rows[0].cells[i]
    set_cell(cell, label, 10, True)
    border_cell(cell)

for row_i in range(1, 5):
    for col_i in range(5):
        cell = table.rows[row_i].cells[col_i]
        set_cell(cell, f'第{row_i}代\nXX\n━━━━━━', 9)
        border_cell(cell)

add_p(doc, '', size=4)
add_p(doc, '注：竖线连接上下代，直系居中，旁支分列两侧。空格填入族人姓名。', '宋体', 9)

doc.add_page_break()

# 行传格式
add_h(doc, '四、行传（齿录）格式', '黑体', 14)

table2 = doc.add_table(rows=2, cols=5)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'
col_h = ['世次', '姓名', '生卒', '配偶', '简历']
col_w = [1.5, 2.0, 5.5, 3.5, 4.5]
for i, h in enumerate(col_h):
    cell = table2.rows[0].cells[i]
    set_cell(cell, h, 10, True)
    cell.width = Cm(col_w[i])
    border_cell(cell)
sample = ['一世', 'XX', '1900年-1980年', '配XX氏', '毕业于XX大学']
for i, text in enumerate(sample):
    cell = table2.rows[1].cells[i]
    set_cell(cell, text, 10)
    cell.width = Cm(col_w[i])
    border_cell(cell)
for _ in range(5):
    row = table2.add_row()
    for i in range(5):
        cell = row.cells[i]
        set_cell(cell, '', 10)
        cell.width = Cm(col_w[i])
        border_cell(cell)

doc.save('C:/Users/Administrator/Downloads/家谱模板_苏式版.docx')
print("OK - 苏式版")
