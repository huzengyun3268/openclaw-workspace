# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def set_font(run, name='宋体', size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_h(doc, text, font='黑体', size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, font, size, True)
    return p

def add_p(doc, text, font='宋体', size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, font, size)
    return p

def set_cell(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, '宋体', size, bold)

def border_cell(cell, color='999999'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        e = OxmlElement(f'w:{b}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)

# 封面
add_p(doc, '', size=8)
for _ in range(4): add_p(doc, '', size=4)
add_h(doc, '《×氏宗谱》', '标宋', 28)
add_p(doc, '欧式世系图版', '宋体', 12)
for _ in range(5): add_p(doc, '', size=6)
add_h(doc, '卷　首', '黑体', 14)
add_p(doc, '一世至五世　五世一图　横排右行', '宋体', 11)
for _ in range(5): add_p(doc, '', size=8)
add_h(doc, '公元二零二六年三月　续修', '宋体', 12)

doc.add_page_break()

# 欧式世系图说明
add_h(doc, '一、欧式世系图说明', '黑体', 14)
add_p(doc, '欧式世系图，又称"横行体"，以五世为一表，世代横行，父子竖线相连，同世并列。其优点在于：支系分明，一目了然；便于查阅代数；适合人口众多的大家族。', '宋体', 11)
add_p(doc, '本谱采用欧式排版，自始祖起，每五世为一图，图内横列五世，纵列各房支系，以竖线相连，清晰展示血脉传承关系。', '宋体', 11)

add_h(doc, '二、阅读方法', '黑体', 14)
rules = [
    ('一', '本谱自右向左阅读，右端为始祖或该表之最高祖。'),
    ('二', '每行横列同一世代，依长幼自左向右排列。'),
    ('三', '父子之间以竖线连接，显示血缘关系。'),
    ('四', '每人名下附注生卒年月、配偶、子女概要。'),
    ('五', '续表接前表时，注明"续前表第×世"。'),
]
for num, r in rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'（{num}）{r}')
    set_font(r1, '宋体', 11)

add_h(doc, '三、世系图（示例）', '黑体', 14)
add_p(doc, '以下为五世一表示例，每格内填入族人姓名及基本信息：', '宋体', 10)

# 欧式世系表
headers = ['第一世', '第二世', '第三世', '第四世', '第五世']
table = doc.add_table(rows=4, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表头
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell(cell, h, 10, True)
    border_cell(cell)

# 示例格子
for row_idx in range(1, 4):
    for col_idx in range(5):
        cell = table.rows[row_idx].cells[col_idx]
        set_cell(cell, f'第{row_idx}代\n姓名\n生卒\n配偶', 9)
        border_cell(cell)

add_p(doc, '', size=4)
add_p(doc, '注：空格处填入族人姓名，字号、生卒、配偶另行记录于行传。', '宋体', 9)

doc.add_page_break()

# 行传格式
add_h(doc, '四、行传（齿录）格式', '黑体', 14)
add_p(doc, '行传详细记录每人的生平事迹，与世系图互为补充。', '宋体', 11)

table2 = doc.add_table(rows=2, cols=5)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'
col_h = ['世次', '姓名', '生卒', '配偶', '简历']
col_w = [1.5, 2.0, 5.5, 3.5, 4.5]
for i, h in enumerate(col_h):
    cell = table2.rows[0].cells[i]
    set_cell(cell, h, 10, True)
    cell.width = Cm(col_w[i])
    border_cell(cell)

sample = ['一世', 'XX', '1900年-1980年', '配XX氏', '毕业于XX大学']
for i, text in enumerate(sample):
    cell = table2.rows[1].cells[i]
    set_cell(cell, text, 10)
    cell.width = Cm(col_w[i])
    border_cell(cell)

for _ in range(5):
    row = table2.add_row()
    for i in range(5):
        cell = row.cells[i]
        set_cell(cell, '', 10)
        cell.width = Cm(col_w[i])
        border_cell(cell)

doc.save('C:/Users/Administrator/Downloads/家谱模板_欧式版.docx')
print("OK - 欧式版")
