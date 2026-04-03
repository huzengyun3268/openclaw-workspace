# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def sf(run, name, size, bold=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def h(doc, text, font='黑体', size=14):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); sf(r, font, size, True)

def para(doc, text, font='宋体', size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p2 = doc.add_paragraph(); p2.alignment = align
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(text); sf(r, font, size)

def bc(cell, color='BBBBBB'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top','left','bottom','right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); tcBorders.append(e)
    tcPr.append(tcBorders)

def sc(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''; p2 = cell.paragraphs[0]; p2.alignment = align
    p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(1)
    r = p2.add_run(text); sf(r, '宋体', size, bold)

# 封面
para(doc, '', '宋体', 8)
for _ in range(3): para(doc, '', '宋体', 4)
h(doc, '《XX氏宗谱》', '标宋', 24)
para(doc, '现代简洁版 · 表格体', '宋体', 11)
for _ in range(4): para(doc, '', '宋体', 6)
h(doc, '公元二零二六年三月　续修', '宋体', 11)
doc.add_page_break()

h(doc, '一、简洁版说明', '黑体', 13)
para(doc, '现代简洁版采用紧凑的表格形式，在一页A4纸上可容纳较多信息。适合信息量较大、需要批量录入的家族使用。', '宋体', 9)

h(doc, '二、世系总表', '黑体', 13)

# 大表格
table = doc.add_table(rows=21, cols=12)
table.alignment = 1  # CENTER
table.style = 'Table Grid'

headers = ['世', '姓名', '字', '号', '生年', '卒年', '寿', '配偶', '子', '女', '学历/职业', '备注']
widths = [0.9, 2.0, 1.5, 1.5, 1.8, 1.8, 0.8, 2.5, 1.5, 1.5, 3.0, 3.5]
for ci, (hdr, w) in enumerate(zip(headers, widths)):
    cell = table.rows[0].cells[ci]
    sc(cell, hdr, 8, True)
    cell.width = Cm(w); bc(cell)

sample = ['一', 'XXX', 'XX', '', '1900', '1980', '80', '配XXX', '2', '1', '大学', '']
for row in table.rows[1:]:
    for ci, (text, w) in enumerate(zip(sample, widths)):
        cell = row.cells[ci]
        sc(cell, text, 8)
        cell.width = Cm(w); bc(cell)

para(doc, '', '宋体', 4)
para(doc, f'注：共20行空格，可自行复制行增加。寿=卒年-生年。学历/职业栏填最高学历及主要职业。', '宋体', 8)

doc.add_page_break()
h(doc, '三、简洁版使用说明', '黑体', 13)
for num, r in [
    ('一', '本表为通用格式，可直接打印使用。'),
    ('二', '每行填一人，复制行可增加行数。'),
    ('三', '学历/职业栏填写最高学历及主要职业。'),
    ('四', '备注栏填写其他需要说明的事项。'),
    ('五', '如需增加列，可在Word中右键表格-插入列。'),
    ('六', '此版本适合电脑录入后打印，也可手写。'),
]:
    p3 = doc.add_paragraph(); p3.paragraph_format.space_before = Pt(3); p3.paragraph_format.space_after = Pt(3)
    r1 = p3.add_run(f'（{num}）{r}'); sf(r1, '宋体', 9)

doc.save('C:/Users/Administrator/Downloads/jiapu_simple.docx')
print("OK")
